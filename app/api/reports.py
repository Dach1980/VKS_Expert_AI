"""Project Expert AI — IOS 3.1 report generation API.

Exports the persisted first-pass check using the structure of the reference
normcontrol report: summary -> consolidated register -> detailed findings ->
conclusion -> normative sources -> corrective-action appendix.
Evidence images saved by the checker are embedded with their red rectangles.
"""
from __future__ import annotations

import io
import json
from datetime import datetime
from pathlib import Path

import fitz
from docx import Document
from docx.shared import Inches
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, KeepTogether, PageBreak
from reportlab.lib import colors

from app.api.documents import DOCUMENTS_ROOT

router = APIRouter(prefix="/api/reports", tags=["reports"])
REPORT_TEMPLATE = "reference_normcontrol_report_ios_3.1"


def _safe_text(value) -> str:
    return str(value or "").strip()


def _document_pdf(document_id: str) -> Path | None:
    root = DOCUMENTS_ROOT / document_id
    if not root.exists():
        return None
    source = root / "source.pdf"
    return source if source.exists() else next((p for p in root.iterdir() if p.is_file() and p.suffix.lower() == ".pdf"), None)


def _saved_evidence_image(finding: dict) -> bytes | None:
    value = _safe_text(finding.get("evidence_image"))
    if value:
        path = Path(value)
        if path.exists() and path.is_file():
            try:
                return path.read_bytes()
            except OSError:
                pass
    return None


def _evidence_image(finding: dict) -> bytes | None:
    """Prefer the exact persisted annotation produced by the first-pass checker."""
    saved = _saved_evidence_image(finding)
    if saved:
        return saved
    document_id = _safe_text(finding.get("docId"))
    page_number = int(finding.get("page") or 0)
    bbox = finding.get("bbox")
    if not document_id or page_number <= 0 or not isinstance(bbox, (list, tuple)) or len(bbox) != 4:
        return None
    pdf_path = _document_pdf(document_id)
    if not pdf_path:
        return None
    try:
        pdf = fitz.open(pdf_path)
        if page_number > len(pdf):
            pdf.close()
            return None
        page = pdf[page_number - 1]
        shape = page.new_shape()
        shape.draw_rect(fitz.Rect(*[float(x) for x in bbox]))
        shape.finish(color=(1, 0, 0), width=2)
        shape.commit()
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
        data = pix.tobytes("png")
        pdf.close()
        return data
    except Exception:
        return None


def _load_saved_report(document_id: str) -> dict | None:
    path = DOCUMENTS_ROOT / document_id / "checking" / "first_pass" / "report.json"
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8-sig"))
    except (OSError, json.JSONDecodeError):
        return None


def _findings(report: dict) -> list[dict]:
    return report.get("results") or report.get("checks") or []


def _source_rows(findings: list[dict]) -> list[str]:
    rows, seen = [], set()
    for finding in findings:
        for source in finding.get("normative_sources") or []:
            document = source.get("document", "—")
            text = f"{document} — версия {source.get('version', '—')}, стр. {source.get('page', '—')}"
            if text not in seen:
                seen.add(text)
                rows.append(text)
    return rows


def _summary(report: dict) -> dict:
    return report.get("summary") or {}


def _build_pdf(report: dict) -> bytes:
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4, rightMargin=14 * mm, leftMargin=14 * mm, topMargin=14 * mm, bottomMargin=14 * mm)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Small", parent=styles["BodyText"], fontSize=8.5, leading=11))
    styles.add(ParagraphStyle(name="Finding", parent=styles["Heading2"], fontSize=12, leading=15, spaceBefore=8, spaceAfter=5))
    styles.add(ParagraphStyle(name="RedFinding", parent=styles["Finding"], textColor=colors.red))

    findings = _findings(report)
    s = _summary(report)
    story = [
        Paragraph("PROJECT EXPERT AI", styles["Title"]),
        Paragraph("РАБОЧИЙ ОТЧЁТ ПО ПРОВЕРКЕ ПРОЕКТНОЙ ДОКУМЕНТАЦИИ", styles["Heading1"]),
        Paragraph("Камеральная проверка представленного комплекта", styles["BodyText"]),
        Spacer(1, 4 * mm),
        Paragraph(f"<b>Документ:</b> {_safe_text(report.get('document_name'))}", styles["BodyText"]),
        Paragraph(f"<b>Контрольная дата:</b> {_safe_text(report.get('checked_at')) or datetime.now().strftime('%d.%m.%Y')}", styles["BodyText"]),
        Paragraph(f"<b>Нормативная база:</b> {_safe_text(report.get('normative_document'))}", styles["BodyText"]),
        Paragraph(f"<b>Действующая версия:</b> {_safe_text(report.get('normative_version'))}", styles["BodyText"]),
        Spacer(1, 6 * mm),
        Paragraph("1. Общая оценка", styles["Heading1"]),
        Paragraph(f"Выполнен первый визуально-нормативный проход по {s.get('pages', 0)} страницам. Выявлено результатов: {s.get('total', 0)}; нарушений: {s.get('violations', 0)}; соответствий: {s.get('compliant', 0)}; не подтверждено: {s.get('unchecked', 0)}.", styles["BodyText"]),
        Spacer(1, 4 * mm),
    ]
    summary_table = Table([["Страниц", "Результатов", "Нарушений", "Соответствий", "Не подтверждено", "Критических"], [s.get("pages", 0), s.get("total", 0), s.get("violations", 0), s.get("compliant", 0), s.get("unchecked", 0), s.get("critical", 0)]], colWidths=[28 * mm] * 6)
    summary_table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), .5, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey), ("ALIGN", (0, 0), (-1, -1), "CENTER"), ("FONTSIZE", (0, 0), (-1, -1), 7.5)]))
    story += [summary_table, Spacer(1, 7 * mm), Paragraph("2. Сводный реестр замечаний", styles["Heading1"])]

    register_data = [["№", "Замечание", "Статус", "Класс", "Первоочередное действие"]]
    for i, f in enumerate(findings, 1):
        status = {"violation": "Подтверждено", "compliant": "Соответствует", "unchecked": "Не подтверждено"}.get(f.get("type"), "Не определено")
        register_data.append([str(i), _safe_text(f.get("title")), status, _safe_text(f.get("severity")), _safe_text(f.get("recommendation")) or "Требуется дополнительная проверка"])
    register = Table(register_data, colWidths=[9 * mm, 55 * mm, 27 * mm, 18 * mm, 61 * mm], repeatRows=1)
    register.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), .4, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey), ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 7), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story += [register, Spacer(1, 7 * mm), Paragraph("3. Подробные результаты проверки", styles["Heading1"])]

    for index, finding in enumerate(findings, 1):
        status = {"violation": "ПОДТВЕРЖДЁННОЕ НЕСООТВЕТСТВИЕ", "compliant": "СООТВЕТСТВИЕ", "unchecked": "СООТВЕТСТВИЕ НЕ ПОДТВЕРЖДЕНО"}.get(finding.get("type"), "РЕЗУЛЬТАТ")
        heading = styles["RedFinding"] if finding.get("type") == "violation" else styles["Finding"]
        story.append(KeepTogether([Paragraph(f"3.{index}. {status}: {_safe_text(finding.get('title')) or 'Результат проверки'}", heading), Paragraph(f"<b>Страница PDF:</b> {_safe_text(finding.get('page')) or '—'} &nbsp;&nbsp; <b>Лист:</b> {_safe_text(finding.get('sheet')) or '—'}", styles["Small"]), Paragraph(f"<b>Класс:</b> {_safe_text(finding.get('severity')) or '—'} &nbsp;&nbsp; <b>Уверенность:</b> {_safe_text(finding.get('confidence')) or '—'}", styles["Small"]), Paragraph(f"<b>Норматив:</b> {_safe_text(finding.get('norm')) or '—'} &nbsp;&nbsp; <b>Пункт:</b> {_safe_text(finding.get('clause')) or '—'}", styles["Small"]), Spacer(1, 2 * mm), Paragraph(f"<b>Основание:</b> {_safe_text(finding.get('description')) or '—'}", styles["BodyText"])]))
        if finding.get("evidence_text"):
            story.append(Paragraph(f"<b>Фрагмент документа:</b> {_safe_text(finding.get('evidence_text'))}", styles["BodyText"]))
        if finding.get("recommendation"):
            story.append(Paragraph(f"<b>Рекомендуемое исправление:</b> {_safe_text(finding.get('recommendation'))}", styles["BodyText"]))
        image_bytes = _evidence_image(finding)
        if image_bytes:
            story += [Spacer(1, 3 * mm), Image(io.BytesIO(image_bytes), width=165 * mm, height=105 * mm), Paragraph("Доказательный фрагмент исходной страницы с красной рамкой.", styles["Small"])]
        story.append(Spacer(1, 6 * mm))

    story.append(PageBreak())
    story.append(Paragraph("4. Заключение", styles["Heading1"]))
    story.append(Paragraph(_safe_text(report.get("conclusion")) or "Рекомендуется устранить подтверждённые замечания до принятия рабочей документации. Результаты первого прохода являются предварительными и подлежат обязательной проверке специалистом.", styles["BodyText"]))
    story.append(Paragraph("5. Нормативные и технические источники проверки", styles["Heading1"]))
    sources = _source_rows(findings) or [f"{_safe_text(report.get('normative_document'))} — действующая версия {_safe_text(report.get('normative_version'))}"]
    for i, source in enumerate(sources, 1):
        story.append(Paragraph(f"{i}. {source}", styles["BodyText"]))
    story.append(Spacer(1, 5 * mm))
    story.append(Paragraph("6. Приложение А. Форма контроля устранения замечаний", styles["Heading1"]))
    violations = [x for x in findings if x.get("type") == "violation"]
    appendix = [["№", "Замечание", "Ответственный", "Срок", "Подтверждающий документ", "Статус"]]
    for i, f in enumerate(violations, 1):
        appendix.append([str(i), _safe_text(f.get("title")), "—", "—", _safe_text(f.get("recommendation")) or "Представить доказательства", "Открыто"])
    if len(appendix) == 1:
        appendix.append(["—", "Подтверждённых замечаний нет", "—", "—", "—", "—"])
    app_table = Table(appendix, colWidths=[9 * mm, 53 * mm, 25 * mm, 20 * mm, 55 * mm, 20 * mm], repeatRows=1)
    app_table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), .4, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey), ("FONTSIZE", (0, 0), (-1, -1), 7), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story.append(app_table)
    doc.build(story)
    return output.getvalue()


def _build_docx(report: dict) -> bytes:
    document = Document()
    findings = _findings(report)
    violations = [x for x in findings if x.get("type") == "violation"]
    s = _summary(report)
    document.add_heading("PROJECT EXPERT AI", level=0)
    document.add_heading("РАБОЧИЙ ОТЧЁТ ПО РЕЗУЛЬТАТАМ ПРОВЕРКИ ПРОЕКТНОЙ ДОКУМЕНТАЦИИ", level=1)
    document.add_paragraph("Камеральная проверка представленного комплекта")
    document.add_paragraph(f"Документ: {_safe_text(report.get('document_name'))}")
    document.add_paragraph(f"Контрольная дата: {_safe_text(report.get('checked_at')) or datetime.now().strftime('%d.%m.%Y')}")
    document.add_paragraph(f"Нормативная база: {_safe_text(report.get('normative_document'))}")
    document.add_paragraph(f"Действующая версия: {_safe_text(report.get('normative_version'))}")
    document.add_heading("1. Общая оценка", level=1)
    document.add_paragraph(f"Проверено страниц: {s.get('pages', 0)}. Результатов: {s.get('total', 0)}. Нарушений: {s.get('violations', 0)}. Соответствий: {s.get('compliant', 0)}. Не подтверждено: {s.get('unchecked', 0)}.")
    document.add_heading("2. Сводный реестр замечаний", level=1)
    table = document.add_table(rows=1, cols=5)
    for i, h in enumerate(["№", "Замечание", "Статус", "Класс", "Первоочередное действие"]): table.cell(0, i).text = h
    for i, f in enumerate(findings, 1):
        cells = table.add_row().cells
        cells[0].text = str(i); cells[1].text = _safe_text(f.get("title")); cells[2].text = {"violation": "Подтверждено", "compliant": "Соответствует", "unchecked": "Не подтверждено"}.get(f.get("type"), "—"); cells[3].text = _safe_text(f.get("severity")); cells[4].text = _safe_text(f.get("recommendation")) or "Требуется дополнительная проверка"
    document.add_heading("3. Подробные результаты проверки", level=1)
    for i, f in enumerate(findings, 1):
        status = {"violation": "ПОДТВЕРЖДЁННОЕ НЕСООТВЕТСТВИЕ", "compliant": "СООТВЕТСТВИЕ", "unchecked": "СООТВЕТСТВИЕ НЕ ПОДТВЕРЖДЕНО"}.get(f.get("type"), "РЕЗУЛЬТАТ")
        document.add_heading(f"3.{i}. {status}: {_safe_text(f.get('title'))}", level=2)
        document.add_paragraph(f"Страница PDF: {_safe_text(f.get('page')) or '—'} | Лист: {_safe_text(f.get('sheet')) or '—'} | Класс: {_safe_text(f.get('severity')) or '—'}")
        document.add_paragraph(f"Норматив: {_safe_text(f.get('norm')) or '—'} | Пункт: {_safe_text(f.get('clause')) or '—'}")
        document.add_paragraph(f"Основание: {_safe_text(f.get('description')) or '—'}")
        if f.get("evidence_text"): document.add_paragraph(f"Фрагмент документа: {_safe_text(f.get('evidence_text'))}")
        if f.get("recommendation"): document.add_paragraph(f"Рекомендуемое исправление: {_safe_text(f.get('recommendation'))}")
        image_bytes = _evidence_image(f)
        if image_bytes:
            document.add_picture(io.BytesIO(image_bytes), width=Inches(6.2)); document.add_paragraph("Доказательный фрагмент исходной страницы с красной рамкой.")
    document.add_heading("4. Заключение", level=1)
    document.add_paragraph(_safe_text(report.get("conclusion")) or "Рекомендуется устранить подтверждённые замечания до принятия рабочей документации. Результаты первого прохода являются предварительными и подлежат обязательной проверке специалистом.")
    document.add_heading("5. Нормативные и технические источники проверки", level=1)
    sources = _source_rows(findings) or [f"{_safe_text(report.get('normative_document'))} — действующая версия {_safe_text(report.get('normative_version'))}"]
    for i, source in enumerate(sources, 1): document.add_paragraph(f"{i}. {source}")
    document.add_heading("6. Приложение А. Форма контроля устранения замечаний", level=1)
    app = document.add_table(rows=1, cols=6)
    for i, h in enumerate(["№", "Замечание", "Ответственный", "Срок", "Подтверждающий документ", "Статус"]): app.cell(0, i).text = h
    for i, f in enumerate(violations, 1):
        cells = app.add_row().cells
        cells[0].text = str(i); cells[1].text = _safe_text(f.get("title")); cells[2].text = "—"; cells[3].text = "—"; cells[4].text = _safe_text(f.get("recommendation")) or "Представить доказательства"; cells[5].text = "Открыто"
    output = io.BytesIO(); document.save(output); return output.getvalue()


@router.get("/{document_id}")
def get_report(document_id: str):
    report = _load_saved_report(document_id)
    if report is None: raise HTTPException(status_code=404, detail="Отчёт для документа ещё не сформирован")
    return {"success": True, "template": REPORT_TEMPLATE, **report}


@router.post("/create/{document_id}")
def create_report(document_id: str):
    report = _load_saved_report(document_id)
    if report is None: raise HTTPException(status_code=409, detail="Сначала выполните проверку документа через /api/checks/{document_id}")
    return {"success": True, "template": REPORT_TEMPLATE, **report}


@router.post("/pdf")
def export_pdf(report: dict):
    try: data = _build_pdf(report)
    except Exception as error: raise HTTPException(status_code=500, detail=f"Не удалось сформировать PDF: {error}") from error
    return StreamingResponse(io.BytesIO(data), media_type="application/pdf", headers={"Content-Disposition": f'attachment; filename="project-expert-ai-check-{report.get("document_id", "report")}.pdf"'})


@router.post("/docx")
def export_docx(report: dict):
    try: data = _build_docx(report)
    except Exception as error: raise HTTPException(status_code=500, detail=f"Не удалось сформировать Word: {error}") from error
    return StreamingResponse(io.BytesIO(data), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f'attachment; filename="project-expert-ai-check-{report.get("document_id", "report")}.docx"'})
