"""IOS 3.1 report builder.

Transforms persisted first-pass report.json into the user-facing normcontrol
report. The checker remains responsible for evidence creation; this layer only
organizes findings and embeds the persisted evidence images.
"""
from __future__ import annotations

import io
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


_FONT_READY = False
_FONT_NAME = "Helvetica"
_BOLD_FONT_NAME = "Helvetica-Bold"


def _register_pdf_font() -> tuple[str, str]:
    """Register a Unicode font, preferring fonts available on Windows."""
    global _FONT_READY, _FONT_NAME, _BOLD_FONT_NAME
    if _FONT_READY:
        return _FONT_NAME, _BOLD_FONT_NAME
    candidates = [
        (Path(r"C:\Windows\Fonts\arial.ttf"), Path(r"C:\Windows\Fonts\arialbd.ttf")),
        (Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"), Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")),
        (Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"), Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf")),
    ]
    for regular, bold in candidates:
        if regular.exists() and bold.exists():
            try:
                pdfmetrics.registerFont(TTFont("ProjectExpertUnicode", str(regular)))
                pdfmetrics.registerFont(TTFont("ProjectExpertUnicodeBold", str(bold)))
                _FONT_NAME = "ProjectExpertUnicode"
                _BOLD_FONT_NAME = "ProjectExpertUnicodeBold"
                break
            except Exception:
                continue
    _FONT_READY = True
    return _FONT_NAME, _BOLD_FONT_NAME


def _text(value) -> str:
    return str(value or "").strip()


def _ptext(value) -> str:
    return escape(_text(value)).replace("\n", "<br/>")


def _findings(report: dict) -> list[dict]:
    return report.get("results") or report.get("checks") or []


def _summary(report: dict) -> dict:
    return report.get("summary") or {}


def _evidence(finding: dict) -> bytes | None:
    value = _text(finding.get("evidence_image"))
    if not value:
        return None
    path = Path(value)
    if not path.is_file():
        return None
    try:
        return path.read_bytes()
    except OSError:
        return None


def _sources(findings: list[dict]) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    for finding in findings:
        for source in finding.get("normative_sources") or []:
            document = _text(source.get("document")) or "—"
            version = _text(source.get("version")) or "—"
            page = _text(source.get("page")) or "—"
            value = f"{document} — версия {version}, стр. {page}"
            if value not in seen:
                seen.add(value)
                result.append(value)
    return result


def build_pdf(report: dict) -> bytes:
    regular, bold = _register_pdf_font()
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4, rightMargin=14 * mm, leftMargin=14 * mm, topMargin=14 * mm, bottomMargin=14 * mm)
    base = getSampleStyleSheet()
    styles = {
        "body": ParagraphStyle("ProjectBody", parent=base["BodyText"], fontName=regular, fontSize=9, leading=12),
        "small": ParagraphStyle("ProjectSmall", parent=base["BodyText"], fontName=regular, fontSize=8, leading=10),
        "title": ParagraphStyle("ProjectTitle", parent=base["Title"], fontName=bold, fontSize=20, leading=24, alignment=1),
        "h1": ParagraphStyle("ProjectH1", parent=base["Heading1"], fontName=bold, fontSize=15, leading=18, spaceBefore=8, spaceAfter=6),
        "h2": ParagraphStyle("ProjectH2", parent=base["Heading2"], fontName=bold, fontSize=11, leading=14, spaceBefore=7, spaceAfter=4),
        "red": ParagraphStyle("ProjectRed", parent=base["Heading2"], fontName=bold, fontSize=11, leading=14, textColor=colors.red, spaceBefore=7, spaceAfter=4),
    }
    findings = _findings(report)
    summary = _summary(report)
    story = [
        Paragraph("PROJECT EXPERT AI", styles["title"]),
        Spacer(1, 4 * mm),
        Paragraph("ОТЧЁТ ПО НОРМОКОНТРОЛЮ ПРОЕКТНОЙ ДОКУМЕНТАЦИИ", styles["h1"]),
        Paragraph("Структура отчёта: reference_normcontrol_report_ios_3.1", styles["body"]),
        Spacer(1, 5 * mm),
        Paragraph(f"<b>Документ:</b> {_ptext(report.get('document_name'))}", styles["body"]),
        Paragraph(f"<b>Дата проверки:</b> {_ptext(report.get('checked_at'))}", styles["body"]),
        Paragraph(f"<b>Нормативная база:</b> {_ptext(report.get('normative_document'))}", styles["body"]),
        Paragraph(f"<b>Действующая версия:</b> {_ptext(report.get('normative_version'))}", styles["body"]),
        Spacer(1, 7 * mm),
        Paragraph("1. Общая оценка", styles["h1"]),
        Paragraph(f"Первый проход выполнен по всему загруженному документу: {summary.get('pages', 0)} страниц. Получено результатов: {summary.get('total', 0)}; подтверждённых несоответствий: {summary.get('violations', 0)}; соответствий: {summary.get('compliant', 0)}; неподтверждённых результатов: {summary.get('unchecked', 0)}.", styles["body"]),
        Spacer(1, 4 * mm),
    ]
    overview = Table([
        ["Страниц", "Результатов", "Несоответствий", "Соответствий", "Не подтверждено", "Критических"],
        [summary.get("pages", 0), summary.get("total", 0), summary.get("violations", 0), summary.get("compliant", 0), summary.get("unchecked", 0), summary.get("critical", 0)],
    ], colWidths=[27 * mm] * 6, repeatRows=1)
    overview.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), .4, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey), ("FONTNAME", (0, 0), (-1, -1), regular), ("FONTNAME", (0, 0), (-1, 0), bold), ("FONTSIZE", (0, 0), (-1, -1), 7), ("ALIGN", (0, 0), (-1, -1), "CENTER")]))
    story += [overview, Spacer(1, 7 * mm), Paragraph("2. Сводный реестр замечаний", styles["h1"])]
    register = [["№", "Замечание", "Статус", "Класс", "Действие"]]
    status_map = {"violation": "Подтверждено", "compliant": "Соответствует", "unchecked": "Не подтверждено"}
    for i, finding in enumerate(findings, 1):
        register.append([str(i), _ptext(finding.get("title")), status_map.get(finding.get("type"), "—"), _ptext(finding.get("severity")), _ptext(finding.get("recommendation")) or "Требуется дополнительная проверка"])
    table = Table(register, colWidths=[8 * mm, 54 * mm, 28 * mm, 18 * mm, 62 * mm], repeatRows=1)
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), .4, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey), ("FONTNAME", (0, 0), (-1, -1), regular), ("FONTNAME", (0, 0), (-1, 0), bold), ("FONTSIZE", (0, 0), (-1, -1), 7), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story += [table, Spacer(1, 7 * mm), Paragraph("3. Подробные результаты проверки", styles["h1"])]
    for i, finding in enumerate(findings, 1):
        violation = finding.get("type") == "violation"
        status = "ПОДТВЕРЖДЁННОЕ НЕСООТВЕТСТВИЕ" if violation else status_map.get(finding.get("type"), "РЕЗУЛЬТАТ")
        heading = styles["red"] if violation else styles["h2"]
        blocks = [
            Paragraph(f"3.{i}. {status}: {_ptext(finding.get('title')) or 'Результат проверки'}", heading),
            Paragraph(f"<b>Страница PDF:</b> {_ptext(finding.get('page')) or '—'} &nbsp;&nbsp; <b>Лист:</b> {_ptext(finding.get('sheet')) or '—'}", styles["small"]),
            Paragraph(f"<b>Норматив:</b> {_ptext(finding.get('norm')) or '—'} &nbsp;&nbsp; <b>Пункт:</b> {_ptext(finding.get('clause')) or '—'}", styles["small"]),
            Paragraph(f"<b>Основание:</b> {_ptext(finding.get('description')) or '—'}", styles["body"]),
        ]
        if finding.get("evidence_text"):
            blocks.append(Paragraph(f"<b>Видимый факт:</b> {_ptext(finding.get('evidence_text'))}", styles["body"]))
        if finding.get("recommendation"):
            blocks.append(Paragraph(f"<b>Рекомендуемое исправление:</b> {_ptext(finding.get('recommendation'))}", styles["body"]))
        story.append(KeepTogether(blocks))
        image = _evidence(finding)
        if image:
            story += [Spacer(1, 2 * mm), Image(io.BytesIO(image), width=165 * mm, height=105 * mm), Paragraph("Evidence: исходная страница с красной рамкой выявленного участка.", styles["small"])]
        story.append(Spacer(1, 5 * mm))
    story += [PageBreak(), Paragraph("4. Заключение", styles["h1"]), Paragraph(_ptext(report.get("conclusion")) or "Подтверждённые несоответствия рекомендуется устранить до выпуска документации. Результаты первого прохода являются инструментом предварительного нормоконтроля и требуют экспертной проверки.", styles["body"]), Paragraph("5. Нормативные источники", styles["h1"])]
    sources = _sources(findings) or [f"{_text(report.get('normative_document'))} — действующая версия {_text(report.get('normative_version'))}"]
    for i, source in enumerate(sources, 1):
        story.append(Paragraph(f"{i}. {_ptext(source)}", styles["body"]))
    story += [Spacer(1, 5 * mm), Paragraph("6. Приложение А. Контроль устранения замечаний", styles["h1"])]
    violations = [f for f in findings if f.get("type") == "violation"]
    appendix = [["№", "Замечание", "Ответственный", "Срок", "Подтверждение", "Статус"]]
    for i, finding in enumerate(violations, 1):
        appendix.append([str(i), _ptext(finding.get("title")), "—", "—", _ptext(finding.get("recommendation")) or "Представить доказательства", "Открыто"])
    if not violations:
        appendix.append(["—", "Подтверждённых замечаний нет", "—", "—", "—", "—"])
    app_table = Table(appendix, colWidths=[8 * mm, 52 * mm, 25 * mm, 18 * mm, 58 * mm, 20 * mm], repeatRows=1)
    app_table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), .4, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey), ("FONTNAME", (0, 0), (-1, -1), regular), ("FONTNAME", (0, 0), (-1, 0), bold), ("FONTSIZE", (0, 0), (-1, -1), 7), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story.append(app_table)
    doc.build(story)
    return output.getvalue()


def build_docx(report: dict) -> bytes:
    document = Document()
    findings = _findings(report)
    summary = _summary(report)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    for style_name in ("Title", "Heading 1", "Heading 2"):
        styles[style_name].font.name = "Arial"
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    title = document.add_heading("PROJECT EXPERT AI", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading("ОТЧЁТ ПО НОРМОКОНТРОЛЮ ПРОЕКТНОЙ ДОКУМЕНТАЦИИ", level=1)
    document.add_paragraph("Структура отчёта: reference_normcontrol_report_ios_3.1")
    for label, value in (("Документ", report.get("document_name")), ("Дата проверки", report.get("checked_at")), ("Нормативная база", report.get("normative_document")), ("Действующая версия", report.get("normative_version"))):
        document.add_paragraph(f"{label}: {_text(value)}")
    document.add_heading("1. Общая оценка", level=1)
    document.add_paragraph(f"Первый проход выполнен по всему документу: {summary.get('pages', 0)} страниц. Результатов: {summary.get('total', 0)}; несоответствий: {summary.get('violations', 0)}; соответствий: {summary.get('compliant', 0)}; не подтверждено: {summary.get('unchecked', 0)}.")
    document.add_heading("2. Сводный реестр замечаний", level=1)
    table = document.add_table(rows=1, cols=5)
    for i, value in enumerate(("№", "Замечание", "Статус", "Класс", "Действие")): table.cell(0, i).text = value
    status_map = {"violation": "Подтверждено", "compliant": "Соответствует", "unchecked": "Не подтверждено"}
    for i, finding in enumerate(findings, 1):
        cells = table.add_row().cells
        values = (str(i), _text(finding.get("title")), status_map.get(finding.get("type"), "—"), _text(finding.get("severity")), _text(finding.get("recommendation")) or "Требуется дополнительная проверка")
        for j, value in enumerate(values): cells[j].text = value
    document.add_heading("3. Подробные результаты проверки", level=1)
    for i, finding in enumerate(findings, 1):
        status = "ПОДТВЕРЖДЁННОЕ НЕСООТВЕТСТВИЕ" if finding.get("type") == "violation" else status_map.get(finding.get("type"), "РЕЗУЛЬТАТ")
        document.add_heading(f"3.{i}. {status}: {_text(finding.get('title')) or 'Результат проверки'}", level=2)
        document.add_paragraph(f"Страница PDF: {_text(finding.get('page')) or '—'} | Лист: {_text(finding.get('sheet')) or '—'}")
        document.add_paragraph(f"Норматив: {_text(finding.get('norm')) or '—'} | Пункт: {_text(finding.get('clause')) or '—'}")
        document.add_paragraph(f"Основание: {_text(finding.get('description')) or '—'}")
        if finding.get("evidence_text"): document.add_paragraph(f"Видимый факт: {_text(finding.get('evidence_text'))}")
        if finding.get("recommendation"): document.add_paragraph(f"Рекомендуемое исправление: {_text(finding.get('recommendation'))}")
        image = _evidence(finding)
        if image:
            document.add_picture(io.BytesIO(image), width=Inches(6.2))
            document.add_paragraph("Evidence: исходная страница с красной рамкой выявленного участка.")
    document.add_heading("4. Заключение", level=1)
    document.add_paragraph(_text(report.get("conclusion")) or "Подтверждённые несоответствия рекомендуется устранить до выпуска документации. Результаты первого прохода требуют экспертной проверки.")
    document.add_heading("5. Нормативные источники", level=1)
    for i, source in enumerate(_sources(findings) or [f"{_text(report.get('normative_document'))} — действующая версия {_text(report.get('normative_version'))}"], 1): document.add_paragraph(f"{i}. {source}")
    document.add_heading("6. Приложение А. Контроль устранения замечаний", level=1)
    violations = [f for f in findings if f.get("type") == "violation"]
    app = document.add_table(rows=1, cols=6)
    for i, value in enumerate(("№", "Замечание", "Ответственный", "Срок", "Подтверждение", "Статус")): app.cell(0, i).text = value
    for i, finding in enumerate(violations, 1):
        cells = app.add_row().cells
        for j, value in enumerate((str(i), _text(finding.get("title")), "—", "—", _text(finding.get("recommendation")) or "Представить доказательства", "Открыто")): cells[j].text = value
    if not violations:
        app.add_row().cells[1].text = "Подтверждённых замечаний нет"
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
